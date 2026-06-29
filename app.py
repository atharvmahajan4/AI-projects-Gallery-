"""
ResumeGenius AI — backend
A small Flask app that:
  1. Extracts text from an uploaded PDF/DOCX (or accepts pasted text)
  2. Sends it to the Gemini API with a resume-grading prompt
  3. Returns a structured JSON report to the frontend dashboard

Run with:  python app.py
Then open: http://localhost:5000
"""

import io
import json
import os
import re

import pdfplumber
import requests
from docx import Document
from dotenv import load_dotenv
from flask import Flask, jsonify, render_template, request
from requests.exceptions import RequestException
from werkzeug.exceptions import HTTPException, RequestEntityTooLarge
from werkzeug.utils import secure_filename

load_dotenv()  # reads .env from the project root (same folder as this file)

GEMINI_API_KEY = os.environ.get("GEMINI_API_KEY", "")
GEMINI_MODEL = os.environ.get("GEMINI_MODEL", "gemini-2.5-flash")
GEMINI_URL = f"https://generativelanguage.googleapis.com/v1beta/models/{GEMINI_MODEL}:generateContent"

app = Flask(__name__)
app.config["MAX_CONTENT_LENGTH"] = 10 * 1024 * 1024  # 10 MB upload cap
# Without this, Flask's debug mode shows its interactive HTML traceback page
# for unexpected errors instead of running our JSON error handler below —
# which breaks the frontend's `await res.json()` call.
app.config["PROPAGATE_EXCEPTIONS"] = False

ALLOWED_EXTENSIONS = {"pdf", "docx"}


# ---------------------------------------------------------------------------
# Text extraction
# ---------------------------------------------------------------------------

def extract_text_from_pdf(file_stream) -> str:
    text_parts = []
    with pdfplumber.open(file_stream) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text() or ""
            text_parts.append(page_text)
    return "\n".join(text_parts).strip()


def extract_text_from_docx(file_stream) -> str:
    document = Document(file_stream)
    paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
    # Tables (skills grids, etc.) are common in resumes — pull those too.
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))
    return "\n".join(paragraphs).strip()


def extract_text(file) -> str:
    filename = secure_filename(file.filename or "")
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    if ext not in ALLOWED_EXTENSIONS:
        raise ValueError("Only .pdf and .docx files are supported.")

    stream = io.BytesIO(file.read())
    try:
        if ext == "pdf":
            return extract_text_from_pdf(stream)
        return extract_text_from_docx(stream)
    except ValueError:
        raise
    except Exception as exc:
        # pdfplumber/python-docx raise their own exception types for
        # corrupted, password-protected, or non-PDF/DOCX-shaped files.
        # Catch them all here so a bad upload returns a clean 400 instead
        # of crashing the request with an unhandled 500.
        raise ValueError(
            "Couldn't read that file — it may be corrupted, password-protected, "
            "or not a real PDF/DOCX. Try pasting the resume text instead."
        ) from exc


# ---------------------------------------------------------------------------
# AI prompt + call
# ---------------------------------------------------------------------------

SYSTEM_PROMPT = """You are an experienced technical recruiter and ATS (Applicant \
Tracking System) specialist. You review resumes for software and tech roles and \
give direct, specific, honest feedback — never generic filler praise or generic \
filler criticism. Every observation must reference something actually present in \
the resume text you were given.

Respond with ONLY a single valid JSON object matching exactly this schema and \
these types — no markdown fences, no commentary before or after it:
IMPORTANT SCORING RULES:

Calculate ATS score between 0 and 100 using:

- Skills Match = 30 marks
- Projects = 20 marks
- Experience = 15 marks
- Education = 10 marks
- ATS Formatting = 10 marks
- Keywords Match = 15 marks

Score Guide:
90-100 = Excellent
75-89 = Good
60-74 = Average
40-59 = Needs Improvement
Below 40 = Poor

Never give random low scores.
A resume with good projects, programming skills, certifications and relevant keywords should normally score between 75 and 90.
Return only an integer ATS score.

{
  "atsScore": number,                  // 0-100, how well an ATS would parse and rank this resume
  "summary": string,                   // <= 60 words, a neutral third-person overview of the candidate
  "missingSkills": [string],           // up to 5 items, short skill names relevant to the resume/role, <= 6 words each
  "strengths": [string],               // up to 4 items, specific to this resume, <= 16 words each
  "weaknesses": [string],              // up to 4 items, specific to this resume, <= 16 words each
  "suggestions": [string],             // up to 5 items, ordered by priority, imperative voice, <= 18 words each
  "interviewQuestions": [string]       // up to 5 full interview questions a recruiter would ask based on THIS resume
}

Keep every string concrete and tied to this specific resume's actual content — \
never write something generic that could apply to any resume."""


def build_user_message(resume_text: str, role: str) -> str:
    parts = ["Here is the candidate's resume text:\n\n" + resume_text.strip()]
    if role:
        parts.append(f"\nTarget role: {role}")
    parts.append(
        "\nAnalyze this resume and return the JSON object described in your "
        "instructions. Be specific to this resume's actual content."
    )
    return "\n".join(parts)


def strip_code_fences(text: str) -> str:
    text = text.strip()
    text = re.sub(r"^```(?:json)?\s*", "", text)
    text = re.sub(r"\s*```$", "", text)
    return text.strip()


def call_gemini(resume_text: str, role: str) -> dict:
    if not GEMINI_API_KEY:
        raise RuntimeError(
            "No GEMINI_API_KEY set. Copy .env.example to .env and add your key."
        )

    payload = {
        # Google's REST API uses camelCase field names. "system_instruction"
        # (snake_case) is NOT a recognized field — sending it gets rejected
        # with a 400 "Invalid JSON payload received. Unknown name
        # \"system_instruction\": Cannot find field." error.
        "systemInstruction": {"parts": [{"text": SYSTEM_PROMPT}]},
        "contents": [
            {"role": "user", "parts": [{"text": build_user_message(resume_text, role)}]}
        ],
        "generationConfig": {
            "temperature": 0.4,
            # 1536 was too tight — a full strengths/weaknesses/suggestions/
            # interview-questions JSON payload could get cut off mid-string,
            # which then failed to parse as JSON. 3072 leaves real headroom.
            "maxOutputTokens": 3072,
            "responseMimeType": "application/json",
        },
    }
    headers = {
        "x-goog-api-key": GEMINI_API_KEY,
        "Content-Type": "application/json",
    }

    try:
        response = requests.post(GEMINI_URL, headers=headers, json=payload, timeout=60)
    except RequestException as exc:
        # DNS failure, no internet, connection refused, timeout, etc.
        raise RuntimeError(f"Couldn't reach the Gemini API: {exc}") from exc

    if response.status_code == 429:
        raise RuntimeError(
            "Gemini rate/quota limit hit: "
            f"{response.text[:300]} — if this says 'PerDay', waiting won't "
            "help until the daily quota resets (midnight Pacific time); if "
            "it says 'PerMinute', wait ~60 seconds and try again."
        )
    if response.status_code != 200:
        raise RuntimeError(f"AI service error ({response.status_code}): {response.text[:300]}")

    try:
        data = response.json()
    except ValueError as exc:
        raise RuntimeError(f"Gemini returned a non-JSON response: {exc}") from exc

    candidates = data.get("candidates") or []
    if not candidates:
        block_reason = data.get("promptFeedback", {}).get("blockReason")
        suffix = f" (blocked: {block_reason})" if block_reason else ""
        raise RuntimeError(f"The AI service returned no result{suffix}.")

    try:
        parts = candidates[0]["content"]["parts"]
        # Skip any "thought" parts — some Gemini model variants can include
        # internal reasoning alongside the answer; we only want the answer.
        raw = "".join(p.get("text", "") for p in parts if not p.get("thought"))
    except (KeyError, IndexError) as exc:
        raise RuntimeError(f"Unexpected response shape from the AI service: {exc}") from exc

    raw = strip_code_fences(raw)
    try:
        return json.loads(raw)
    except json.JSONDecodeError as exc:
        raise RuntimeError(f"Couldn't parse the AI's response as JSON: {exc}") from exc


# ---------------------------------------------------------------------------
# Routes
# ---------------------------------------------------------------------------

@app.route("/")
def index():
    return render_template("index.html")


@app.route("/api/analyze", methods=["POST"])
def analyze():
    role = (request.form.get("role") or "").strip()
    pasted_text = (request.form.get("resume_text") or "").strip()

    resume_text = pasted_text
    uploaded_file = request.files.get("resume_file")

    if uploaded_file and uploaded_file.filename:
        try:
            extracted = extract_text(uploaded_file)
        except ValueError as exc:
            return jsonify({"error": str(exc)}), 400
        if extracted:
            resume_text = extracted

    if not resume_text or len(resume_text) < 40:
        return jsonify({
            "error": "Couldn't find enough resume text. Upload a PDF/DOCX or paste the text."
        }), 400

    if len(resume_text) > 18000:
        resume_text = resume_text[:18000]

    try:
        report = call_gemini(resume_text, role)
    except RuntimeError as exc:
        return jsonify({"error": str(exc)}), 502

    return jsonify(report)


@app.route("/api/health")
def health():
    return jsonify({"ok": True, "model": GEMINI_MODEL, "key_configured": bool(GEMINI_API_KEY)})


# ---------------------------------------------------------------------------
# Error handlers — the frontend always does `await res.json()`, so every
# response (success or failure) needs to actually be JSON, never Flask's
# default HTML error page.
# ---------------------------------------------------------------------------

@app.errorhandler(RequestEntityTooLarge)
def handle_file_too_large(_exc):
    return jsonify({
        "error": "That file is larger than the 10 MB limit. Try a smaller file or paste the text instead."
    }), 413


@app.errorhandler(404)
def handle_not_found(_exc):
    return jsonify({"error": "Not found."}), 404


@app.errorhandler(HTTPException)
def handle_http_exception(exc):
    return jsonify({"error": exc.description or exc.name}), exc.code


@app.errorhandler(Exception)
def handle_unexpected_error(exc):
    app.logger.exception("Unexpected error")
    return jsonify({"error": "An unexpected server error occurred. Please try again."}), 500


if __name__ == "__main__":
    app.run(debug=True, port=5000)
